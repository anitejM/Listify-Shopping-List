from PyQt6.QtWidgets import QApplication, QWidget, QVBoxLayout, QLabel, QLineEdit, QPushButton, QListWidget, QMessageBox, QFileDialog
from PyQt6.QtGui import QFont, QClipboard
from docx import Document
from docx.shared import Pt
import sys

class ShoppingListApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Listify - Smart Shopping List")
        self.setGeometry(100, 100, 400, 500)
        self.total_cost = 0.0
        
        layout = QVBoxLayout()
        
        self.title_label = QLabel("Listify")
        self.title_label.setFont(QFont("Arial", 24, QFont.Weight.Bold))
        layout.addWidget(self.title_label)
        
        self.subtitle_label = QLabel("Made with ♥ by Anitej Mishra")
        self.subtitle_label.setFont(QFont("Arial", 10))
        layout.addWidget(self.subtitle_label)
        
        self.item_input = QLineEdit()
        self.item_input.setPlaceholderText("Enter item name")
        layout.addWidget(self.item_input)
        
        self.quantity_input = QLineEdit()
        self.quantity_input.setPlaceholderText("Enter quantity")
        layout.addWidget(self.quantity_input)
        
        self.cost_input = QLineEdit()
        self.cost_input.setPlaceholderText("Enter cost per unit")
        layout.addWidget(self.cost_input)
        
        self.add_button = QPushButton("Add Item")
        self.add_button.clicked.connect(self.add_item)
        layout.addWidget(self.add_button)
        
        self.list_widget = QListWidget()
        layout.addWidget(self.list_widget)
        
        self.total_label = QLabel("Total Cost: ₹0.00")
        layout.addWidget(self.total_label)
        
        self.clear_button = QPushButton("Clear List")
        self.clear_button.clicked.connect(self.clear_list)
        layout.addWidget(self.clear_button)
        
        self.copy_button = QPushButton("Copy to Clipboard")
        self.copy_button.clicked.connect(self.copy_to_clipboard)
        layout.addWidget(self.copy_button)
        
        self.save_button = QPushButton("Save as Word Document")
        self.save_button.clicked.connect(self.save_to_doc)
        layout.addWidget(self.save_button)
        
        self.setLayout(layout)
    
    def add_item(self):
        item = self.item_input.text().strip()
        quantity = self.quantity_input.text().strip()
        cost = self.cost_input.text().strip()
        
        if not item or not quantity.isdigit() or not cost.replace('.', '', 1).isdigit():
            QMessageBox.critical(self, "Error", "Please enter valid details!")
            return
        
        quantity = int(quantity)
        cost = float(cost)
        total_item_cost = quantity * cost
        self.total_cost += total_item_cost
        
        self.list_widget.addItem(f"{item} - Qty: {quantity} - Cost: ₹{total_item_cost:.2f}")
        self.total_label.setText(f"Total Cost: ₹{self.total_cost:.2f}")
    
    def clear_list(self):
        self.list_widget.clear()
        self.total_cost = 0.0
        self.total_label.setText("Total Cost: ₹0.00")
    
    def copy_to_clipboard(self):
        text = "Shopping List - Listify\n-----------------------------\n"
        for i in range(self.list_widget.count()):
            text += self.list_widget.item(i).text() + "\n"
        text += f"\nTotal Cost: ₹{self.total_cost:.2f}"
        
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        QMessageBox.information(self, "Copied", "Shopping list copied to clipboard!")
    
    def save_to_doc(self):
        file_name, _ = QFileDialog.getSaveFileName(self, "Save File", "", "Word Documents (*.docx)")
        
        if file_name:
            doc = Document()
            doc.add_heading("Shopping List - Listify", level=1)
            doc.add_paragraph("-----------------------------")
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Item'
            hdr_cells[1].text = 'Quantity'
            hdr_cells[2].text = 'Total Cost'
            
            for i in range(self.list_widget.count()):
                entry = self.list_widget.item(i).text().split(' - ')
                row_cells = table.add_row().cells
                row_cells[0].text = entry[0]
                row_cells[1].text = entry[1].split(': ')[1]
                row_cells[2].text = entry[2].split('₹')[1]
            
            doc.add_paragraph(f"\nTotal Cost: ₹{self.total_cost:.2f}")
            doc.save(file_name)
            QMessageBox.information(self, "Saved", "Shopping list saved successfully!")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ShoppingListApp()
    window.show()
    sys.exit(app.exec())
